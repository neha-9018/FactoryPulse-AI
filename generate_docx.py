import sys
import subprocess
import os

def install_and_generate():
    print("Checking for python-docx library...")
    try:
        import docx
    except ImportError:
        print("Installing python-docx...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        import docx

    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    print("Generating Microsoft Word Document (.docx)...")
    doc = Document()

    # Base Styles Setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Document Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Smart Factory Industry 4.0 Dashboard")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 23, 42)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Centralized Monitoring System Fusing SCADA Telemetry, ML Failure Forecasting, and CV Quality Audits")
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph().paragraph_format.space_after = Pt(20)

    # Section 1
    doc.add_heading("1. Project Overview & Objective", level=1)
    p = doc.add_paragraph(
        "The objective of this project is to develop a centralized Industry 4.0 analytics platform that automates shop floor control "
        "and provides visibility into operational efficiency. By combining IoT sensor pipelines, Machine Learning algorithms, "
        "and Computer Vision inspections, the platform minimizes unplanned downtime and optimizes manufacturing yields."
    )
    p.paragraph_format.space_after = Pt(12)

    # Section 2
    doc.add_heading("2. Machine Learning Architecture (ML Pipeline)", level=1)
    p = doc.add_paragraph(
        "The predictive diagnostics pipeline continuously monitors physical parameters like temperature, current draw, and vibration. "
        "Telemetry datasets are normalized using standard scalers and processed by an ML classifier."
    )
    p.paragraph_format.space_after = Pt(12)

    doc.add_heading("A. Dynamic Health Index Formula", level=2)
    formula = doc.add_paragraph()
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = formula.add_run("Health Score = 40% (Vibration) + 30% (Temperature) + 20% (Current) + 10% (Reject Rate)")
    run.font.bold = True
    run.font.size = Pt(12)

    doc.add_heading("B. Health Status Boundaries", level=2)
    doc.add_paragraph("• Healthy (>80%): Safe operational levels; standard schedules active.")
    doc.add_paragraph("• Warning (60% - 80%): Minor wear anomalies; inspection advised.")
    doc.add_paragraph("• Critical (<60%): High-risk operational failure; auto-stop triggered.")

    # Section 3
    doc.add_heading("3. Computer Vision Inspection", level=1)
    doc.add_paragraph(
        "The conveyor check station utilizes grayscaled camera checks and boundary pixel mapping models to identify parts status instantly."
    )
    doc.add_paragraph("• LABEL_MISSING: Fails boundary checks for sticker packaging.")
    doc.add_paragraph("• SURFACE_CRACK: OpenCV contour edge checks.")
    doc.add_paragraph("• DAMAGE: Discrepancies in target surface area bounds.")
    doc.add_paragraph("• WRONG_COLOR / DIMENSION: Dimension deviations.")

    # Section 4
    doc.add_heading("4. Role-Based Access Control (RBAC) Matrix", level=1)
    
    # Create Table
    table = doc.add_table(rows=8, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ["Module", "Admin", "Manager", "Engineer", "Operator"]
    hdr_cells = table.rows[0].cells
    for idx, name in enumerate(headers):
        hdr_cells[idx].text = name
        hdr_cells[idx].paragraphs[0].runs[0].font.bold = True
        
    matrix_data = [
        ["Executive View", "✅ Full Access", "✅ Full Access", "👁️ View Only", "❌ Denied"],
        ["Production Analytics", "✅ Full Access", "✅ Full Access", "✅ Full Access", "👁️ Shift Locked"],
        ["Quality Control", "✅ Full Access", "✅ Full Access", "✅ Full Access", "👁️ View + Defect check"],
        ["Maintenance AI", "✅ Full Access", "✅ Full Access", "✅ Full Access", "✅ Full Access"],
        ["Custom Analytics", "✅ Full Access", "✅ Full Access", "✅ Full Access", "👁️ Limited View"],
        ["AI Assistant", "✅ Full Access", "✅ Full Access", "✅ Full Access", "✅ Full Access"],
        ["Corporate Reports", "✅ Full Access", "✅ Full Access", "👁️ Export (Health)", "❌ Denied"]
    ]
    
    for row_idx, row_data in enumerate(matrix_data):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = text

    # Section 5
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    doc.add_heading("5. Auto-Stop SCADA Request Triggers", level=1)
    doc.add_paragraph("The dashboard implements automated safe stop requests to SCADA under two conditions:")
    doc.add_paragraph("1. Critical Health Drop: If the calculated ML health score falls below 60%, the backend triggers a safe shutdown override, setting status to OFFLINE and dispatching a PLC Stop alert.")
    doc.add_paragraph("2. Defects Rate Override: If 3 consecutive products fail the OpenCV inspections, the PLC auto-stops the conveyor spindle line to inspect the machine clamps.")

    # Save Document
    filename = "smart_factory_project_report.docx"
    doc.save(filename)
    print(f"[SUCCESS] Microsoft Word Document compiled successfully as {filename}!")

if __name__ == "__main__":
    install_and_generate()
